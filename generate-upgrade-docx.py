#!/usr/bin/env python3
"""Generate the AgentStudio Upgrade Plan DOCX with all specialist reviews."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with headers and rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2D3748')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F7FAFC')

    return table


def add_code_block(doc, code, language='typescript'):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(45, 55, 72)


def add_section_header(doc, text, level=2):
    """Add a section header with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(26, 32, 44)
    return heading


def add_specialist_badge(doc, name, icon):
    """Add a specialist review badge."""
    p = doc.add_paragraph()
    run = p.add_run(f'{icon} {name} Review')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(49, 130, 206)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_priority_item(doc, priority, text):
    """Add a priority-tagged item."""
    p = doc.add_paragraph(style='List Bullet')
    if priority == 'P0':
        run = p.add_run('[P0 CRITICAL] ')
        run.bold = True
        run.font.color.rgb = RGBColor(197, 48, 48)
    elif priority == 'P1':
        run = p.add_run('[P1 HIGH] ')
        run.bold = True
        run.font.color.rgb = RGBColor(214, 158, 46)
    elif priority == 'P2':
        run = p.add_run('[P2 MEDIUM] ')
        run.bold = True
        run.font.color.rgb = RGBColor(49, 130, 206)
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)


def build_document():
    doc = Document()

    # =====================================================================
    # STYLES
    # =====================================================================
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(26, 32, 44)

    # =====================================================================
    # COVER PAGE
    # =====================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AgentStudio')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(49, 130, 206)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Multi-Agent Intelligence Layer')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(74, 85, 104)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Comprehensive Upgrade Plan with Specialist Reviews')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(113, 128, 150)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'Generated: {datetime.date.today().strftime("%B %d, %Y")}\n')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(160, 174, 192)
    run = meta.add_run('Reviews: Electron Architect | UX/UI Specialist | Agentic Architect | DB Architect\n')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(160, 174, 192)
    run = meta.add_run('Reference Projects: DevTeam | wshobson/agents | Multi-Agent Squad')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(160, 174, 192)

    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Critical Issues (Cross-Cutting)',
        '3. Phase 1: Complexity Scoring & Model Routing',
        '4. Phase 2: Task Loop with Quality Gates',
        '5. Phase 3: Anti-Abandonment Detection',
        '6. Phase 4: File-Based Agent Communication Chain',
        '7. Phase 5: Cost Tracking Dashboard',
        '8. Phase 6: Human Checkpoint UI',
        '9. Phase 7: Progressive Skill Loading',
        '10. Phase 8: Scope Enforcement Layer',
        '11. Phase 9: Declarative Hooks System',
        '12. Phase 10: Deep Agent Personas & Bug Council',
        '13. Database Migration Guide (Complete SQL)',
        '14. UX/UI Component Specifications',
        '15. Implementation Order & Sprint Plan',
        'Appendix A: File Creation/Modification Index',
        'Appendix B: IPC Channel Registry',
        'Appendix C: Reference Repo Quick Lookup',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(49, 130, 206)

    doc.add_page_break()

    # =====================================================================
    # 1. EXECUTIVE SUMMARY
    # =====================================================================
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This document presents a phased implementation plan to upgrade AgentStudio from a well-built '
        'Electron desktop application into a production-grade multi-agent orchestration system. The plan '
        'integrates the best patterns from three open-source reference projects:'
    )

    refs = [
        ('DevTeam', 'Model escalation, quality gates, anti-abandonment, scope enforcement, cost tracking, Bug Council'),
        ('wshobson/agents', 'Progressive skill loading, file-based agent communication, TDD workflows, phase checkpoints'),
        ('Multi-Agent Squad', 'Declarative hooks (TOML), deep agent personas, human checkpoint patterns, git worktree orchestration'),
    ]
    for name, desc in refs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph(
        'Four specialist agents have reviewed this plan and contributed refinements:'
    )

    specialists = [
        ('Electron Architect', 'IPC patterns, process boundaries, security, child process management'),
        ('UX/UI Specialist', '13 new components with full Tailwind CSS 4 specs, wireframes, accessibility'),
        ('Agentic Architect', 'Multi-agent orchestration, prompt engineering, Claude CLI compatibility, context window management'),
        ('DB Architect', 'Complete SQL migrations, 6 new tables, 14 new indexes, 4 repository classes'),
    ]
    for name, desc in specialists:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()

    doc.add_heading('Current Assessment', level=2)
    doc.add_paragraph(
        'AgentStudio scores 7.5/10 on code quality. It has excellent Electron architecture (contextIsolation, '
        'IPC validation, Zustand state management) but lacks intelligence in agent coordination. '
        'The core gap: "smart infrastructure, dumb orchestration" \u2014 great desktop app, but specialists '
        'all use the same model, have no quality gates, can abandon tasks silently, and have no scope restrictions.'
    )

    doc.add_heading('Expected Outcomes', level=2)
    outcomes = [
        '40-60% API cost savings from model routing (Phase 1)',
        '80%+ task success rate from quality gates + model escalation (Phase 2)',
        'Zero silent abandonments from anti-abandonment detection (Phase 3)',
        'Full audit trail from file-based communication (Phase 4)',
        'User cost visibility and budget control (Phase 5)',
        'Human control at critical decision points (Phase 6)',
        'Reduced token usage from progressive skill loading (Phase 7)',
        'Zero scope violations from enforcement layer (Phase 8)',
        'User-customizable automation via declarative hooks (Phase 9)',
        'Expert-level diagnostics from Bug Council (Phase 10)',
    ]
    for o in outcomes:
        doc.add_paragraph(o, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # 2. CRITICAL ISSUES (CROSS-CUTTING)
    # =====================================================================
    doc.add_heading('2. Critical Issues (Cross-Cutting)', level=1)
    doc.add_paragraph(
        'The following issues were identified by specialist reviews and affect multiple phases. '
        'These must be addressed before or during implementation.'
    )

    doc.add_heading('P0 \u2014 Must Fix Before Implementation', level=2)

    add_priority_item(doc, 'P0',
        'Specialist stdin is \'ignore\' \u2014 Cannot inject re-engagement prompts into running specialists. '
        'specialist-pool.service.ts spawns Claude CLI with stdin: \'ignore\'. Phase 3 (anti-abandonment) '
        'cannot write to specialist stdin. SOLUTION: For specialists, spawn a NEW process with fix context '
        'instead of writing to stdin. Re-engagement via stdin only works for the generalist (long-lived session).')

    add_priority_item(doc, 'P0',
        'Scope validation happens after worktree merge \u2014 specialist-pool.service.ts auto-merges worktrees '
        'on specialist completion. If scope violations exist, they\'re already merged. SOLUTION: Defer merge '
        'until after scope validation (Phase 8) + checkpoint approval (Phase 6). Add a \'validated\' flag to '
        'worktree state.')

    add_priority_item(doc, 'P0',
        'Worktree not reused across task loop iterations \u2014 Each specialist spawn creates a fresh worktree, '
        'losing all prior work. SOLUTION: Pass the same worktree path to all iterations of the same task. '
        'Only create a new worktree on the first iteration.')

    add_priority_item(doc, 'P0',
        'Wrong stdin message format for generalist re-engagement \u2014 The generalist expects NDJSON format: '
        '{ type: "user", message: { role: "user", content: prompt } }. Phase 3 code must match this format exactly.')

    doc.add_heading('P1 \u2014 High Priority', level=2)

    add_priority_item(doc, 'P1',
        'Keyword-based complexity scoring is unreliable on 1-2 sentence task descriptions. '
        'SOLUTION: Use a single batched haiku LLM call to score all tasks at once (cheaper and more accurate). '
        'Fall back to keyword scoring only if the LLM call fails.')

    add_priority_item(doc, 'P1',
        'Fix context grows unbounded across task loop iterations. After 5+ iterations, the accumulated fix '
        'context may exceed the context window. SOLUTION: Cap fix context at ~1000 tokens. Keep only the '
        'last 2 iteration results. Summarize older iterations.')

    add_priority_item(doc, 'P1',
        'Specialist token tracking returns 0 for all specialists (line ~350 in specialist-pool.service.ts). '
        'Token data comes from agent-base.service.ts parsing message_start/message_delta events, but specialists '
        'may not emit these in -p mode. SOLUTION: Verify token extraction works in print mode, or parse usage '
        'from the final message_stop event.')

    add_priority_item(doc, 'P1',
        'No global process pool / concurrency limiter. Multiple phases spawn Claude CLI processes. Without a '
        'limiter, the system could spawn 10+ simultaneous processes. SOLUTION: Create a ProcessPoolService '
        'with configurable max concurrency (default: 4). All specialist spawning goes through this pool.')

    doc.add_heading('P2 \u2014 Medium Priority', level=2)

    add_priority_item(doc, 'P2',
        'Database migration system is informal. Currently uses inline ALTER TABLE wrapped in try/catch. '
        'Six new tables and 8 column additions across phases need a proper migration system. '
        'SOLUTION: Add a schema_version table and sequential migration functions.')

    add_priority_item(doc, 'P2',
        'Decomposition prompt needs richer metadata. Current DECOMPOSITION_SYSTEM_PROMPT outputs minimal task '
        'objects. Upgrade to emit: estimatedComplexity, taskType, allowedPatterns, riskFlags, estimatedFiles.')

    add_priority_item(doc, 'P2',
        'IPC event flooding risk. Phases 2, 3, 5 add real-time events (loop iterations, gate results, cost updates). '
        'SOLUTION: Throttle UI updates to 200ms minimum interval. Batch multiple events into single IPC messages.')

    doc.add_page_break()

    # =====================================================================
    # PHASE 1
    # =====================================================================
    doc.add_heading('3. Phase 1: Complexity Scoring & Model Routing', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Stop using a single model for all specialists. Score every task by complexity (0-14) and route to '
        'the cheapest capable model. This alone saves 40-60% on API costs.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Complexity scoring (0-14)', 'DevTeam', 'agents/orchestration/task-loop.md'],
            ['Model tier config', 'DevTeam', '.devteam/task-loop-config.yaml'],
            ['Tier assignments (112 agents)', 'wshobson/agents', 'docs/agents.md'],
        ])

    doc.add_heading('Current State', level=2)
    doc.add_paragraph(
        'constants.ts defines ACTIVATION_MODEL_ID (sonnet) and BRAIN_FEED_MODEL_ID (haiku) but these are only '
        'used for workspace activation and brain summarization. Orchestrator and specialist pool spawn Claude CLI '
        'without any --model flag override. All specialists inherit the CLI default model.'
    )

    doc.add_heading('What to Build', level=2)

    doc.add_heading('1.1 \u2014 Complexity Scoring Service', level=3)
    doc.add_paragraph('Create: src/main/services/complexity-scorer.service.ts')
    add_code_block(doc, '''// Scoring dimensions (0-14 total):
//   Files affected:     0-3 pts  (1 file=0, 2-3=1, 4-6=2, 7+=3)
//   Estimated lines:    0-3 pts  (<50=0, 50-150=1, 150-300=2, 300+=3)
//   New dependencies:   0-2 pts  (0=0, 1-2=1, 3+=2)
//   Task type:          0-3 pts  (docs=0, test=1, impl=2, arch=3)
//   Risk flags:         0-3 pts  (1pt each: security, external, breaking)
//
// Tier mapping:
//   0-4:   Simple   -> haiku
//   5-8:   Moderate -> sonnet
//   9-14:  Complex  -> opus

export interface ComplexityScore {
  filesAffected: number       // 0-3
  estimatedLines: number      // 0-3
  newDependencies: number     // 0-2
  taskType: number            // 0-3
  riskFlags: number           // 0-3
  total: number               // 0-14
  tier: 'simple' | 'moderate' | 'complex'
  model: 'haiku' | 'sonnet' | 'opus'
}''')

    doc.add_heading('1.2 \u2014 Extend DecomposedTask Type', level=3)
    doc.add_paragraph('In src/shared/types.ts, add new fields:')
    add_code_block(doc, '''export interface DecomposedTask {
  // ...existing fields...
  complexity?: ComplexityScore
  model?: 'haiku' | 'sonnet' | 'opus'
  maxRetries?: number
}''')

    doc.add_heading('1.3 \u2014 Model Routing in Specialist Pool', level=3)
    doc.add_paragraph('In specialist-pool.service.ts, modify spawnSpecialist() to pass --model flag:')
    add_code_block(doc, '''const model = task.model || this.getModelForComplexity(task.complexity?.total || 5)
const args = ['--yes', '-p', task.description, '--model', model, ...]

private getModelForComplexity(score: number): string {
  if (score <= 4) return 'claude-haiku-4-20250414'
  if (score <= 8) return 'claude-sonnet-4-20250514'
  return 'claude-opus-4-20250514'
}''')

    doc.add_heading('1.4 \u2014 Workspace-Level Cost Preference', level=3)
    add_code_block(doc, '''export type CostPreference = 'economy' | 'balanced' | 'power'
// economy: always start with haiku, escalate slower
// balanced: use complexity scoring (default)
// power: always use opus''')

    # Specialist Reviews for Phase 1
    doc.add_heading('Specialist Reviews \u2014 Phase 1', level=2)

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'RECOMMENDATION: Keyword-based scoring on 1-2 sentence task descriptions is unreliable. '
        'Use a single batched haiku LLM call to score ALL decomposed tasks at once. This costs ~$0.001 per '
        'batch and provides accurate structured JSON scores. Fall back to keyword scoring only if the LLM call fails.'
    )
    doc.add_paragraph(
        'Also upgrade DECOMPOSITION_SYSTEM_PROMPT to emit richer metadata: estimatedComplexity, taskType, '
        'allowedPatterns (for Phase 8 scope), riskFlags. This reduces the need for a separate scoring step.'
    )

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'The --model flag must be validated against known model IDs before passing to spawn(). '
        'Add model ID constants to constants.ts: MODEL_IDS = { haiku: "claude-haiku-4-20250414", '
        'sonnet: "claude-sonnet-4-20250514", opus: "claude-opus-4-20250514" } as const.'
    )

    add_specialist_badge(doc, 'UX/UI Specialist', '\U0001F3A8')
    doc.add_paragraph('New component: CostPreferenceSelector')
    doc.add_paragraph(
        'A 3-segment toggle control (Economy / Balanced / Power) in Workspace Settings. Uses Tailwind CSS 4 with '
        'emerald-500 for Economy, blue-500 for Balanced, purple-500 for Power. Shows estimated cost impact: '
        '"~40% savings" / "Recommended" / "Max capability, ~3x cost". Stores preference in workspace settings_json column.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    doc.add_paragraph('Schema additions:')
    add_code_block(doc, '''-- Phase 1: Add to agent_sessions
ALTER TABLE agent_sessions ADD COLUMN complexity_score INTEGER;
ALTER TABLE agent_sessions ADD COLUMN complexity_tier TEXT;
ALTER TABLE agent_sessions ADD COLUMN model_used TEXT;
ALTER TABLE agent_sessions ADD COLUMN tokens_input INTEGER DEFAULT 0;
ALTER TABLE agent_sessions ADD COLUMN tokens_output INTEGER DEFAULT 0;

-- NOTE: Split existing token_usage (single TEXT field with JSON)
-- into tokens_input + tokens_output (INTEGER) for proper aggregation.

-- Model pricing reference table
CREATE TABLE IF NOT EXISTS model_pricing (
  model_id TEXT PRIMARY KEY,
  input_per_1m_cents INTEGER NOT NULL,
  output_per_1m_cents INTEGER NOT NULL,
  effective_date TEXT NOT NULL DEFAULT (datetime('now'))
);

INSERT OR IGNORE INTO model_pricing VALUES
  ('claude-haiku-4-20250414', 80, 400, datetime('now')),
  ('claude-sonnet-4-20250514', 300, 1500, datetime('now')),
  ('claude-opus-4-20250514', 1500, 7500, datetime('now'));''')

    doc.add_heading('Files to Create/Modify', level=2)
    add_styled_table(doc,
        ['Action', 'File'],
        [
            ['CREATE', 'src/main/services/complexity-scorer.service.ts'],
            ['MODIFY', 'src/shared/types.ts \u2014 extend DecomposedTask, add ComplexityScore, CostPreference'],
            ['MODIFY', 'src/shared/constants.ts \u2014 add MODEL_IDS constants per tier'],
            ['MODIFY', 'src/main/services/orchestrator.service.ts \u2014 wire scoring after decompose'],
            ['MODIFY', 'src/main/services/specialist-pool.service.ts \u2014 pass --model flag'],
            ['MODIFY', 'src/main/db/schema.sql \u2014 add columns + model_pricing table'],
            ['MODIFY', 'src/main/db/repositories/agent-session.repository.ts \u2014 persist new fields'],
            ['CREATE', 'src/renderer/src/components/settings/CostPreferenceSelector.tsx'],
        ])

    doc.add_heading('Acceptance Criteria', level=2)
    criteria = [
        'Every specialist task gets a complexity score before execution',
        'Specialists spawn with the correct --model flag based on score',
        'Users can set a workspace-level cost preference (economy/balanced/power)',
        'agent_sessions table records which model was used and the complexity score',
        'Simple tasks (docs, config changes) consistently route to haiku',
        'Complex tasks (architecture, security) consistently route to opus',
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # PHASE 2
    # =====================================================================
    doc.add_heading('4. Phase 2: Task Loop with Quality Gates', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Wrap every specialist execution in an iterative loop: Execute -> Validate -> Pass? Done. '
        'Fail? Fix -> Escalate model -> Retry. No specialist completes until quality gates pass.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Task loop architecture', 'DevTeam', 'agents/orchestration/task-loop.md'],
            ['Quality gate commands', 'DevTeam', 'agents/orchestration/quality-gate-enforcer.md'],
            ['Iteration limits', 'DevTeam', '.devteam/task-loop-config.yaml'],
            ['PostToolUse auto-test', 'Multi-Agent Squad', '.claude/hooks/enterprise-workflow.toml'],
            ['TDD red-green-refactor', 'wshobson/agents', 'plugins/conductor/commands/implement.md'],
        ])

    doc.add_heading('Current State', level=2)
    doc.add_paragraph(
        'specialist-pool.service.ts has retry logic (maxRetries=2, exponential backoff) but retries use the '
        'SAME model and approach. A specialist "succeeds" simply by exiting with code 0. No post-execution '
        'test/lint/typecheck validation exists. No model escalation on retry.'
    )

    doc.add_heading('What to Build', level=2)

    doc.add_heading('2.1 \u2014 Quality Gate Service', level=3)
    doc.add_paragraph('Create: src/main/services/quality-gate.service.ts')
    add_code_block(doc, '''export interface GateResult {
  gate: 'tests' | 'typecheck' | 'lint' | 'security'
  passed: boolean
  output: string
  command: string
  durationMs: number
  exitCode: number
}

export interface QualityGateReport {
  allPassed: boolean
  gates: GateResult[]
  failedGates: string[]
  timestamp: string
}

export class QualityGateService {
  async detectGates(workspacePath: string): Promise<string[]>
  async runGates(workspacePath: string): Promise<QualityGateReport>
}''')

    doc.add_heading('2.2 \u2014 Task Loop Wrapper', level=3)
    doc.add_paragraph('Create: src/main/services/task-loop.service.ts')
    doc.add_paragraph('Loop architecture:')
    loop_steps = [
        '1. Execute specialist with current model',
        '2. Wait for specialist to complete',
        '3. If task is non-code (docs, config): mark COMPLETE, break',
        '4. Run quality gates via QualityGateService',
        '5. If all gates pass: mark COMPLETE, break',
        '6. If gates fail:',
        '   a. Increment consecutiveFailures',
        '   b. Check stuck detection (same test failing 3x -> force escalate)',
        '   c. If consecutiveFailures >= threshold: haiku->sonnet->opus',
        '   d. Build fix context from gate output',
        '   e. Append fix context to task description for next iteration',
        '7. Emit "iteration" event with current state',
    ]
    for step in loop_steps:
        doc.add_paragraph(step, style='List Bullet')

    add_code_block(doc, '''export interface TaskLoopConfig {
  maxIterations: number           // default: 10
  escalationThreshold: number     // consecutive failures before escalate (default: 2)
  opusMaxFailures: number         // failures at opus before giving up (default: 3)
  runQualityGates: boolean        // can disable for non-code tasks
  gateTimeout: number             // ms, default: 120000
}

export interface TaskLoopState {
  iteration: number
  currentModel: string
  consecutiveFailures: number
  failureHistory: Array<{
    iteration: number; model: string;
    failedGates: string[]; errorSummary: string;
  }>
  status: 'running' | 'passed' | 'failed' | 'escalated' | 'max_iterations'
}''')

    doc.add_heading('2.3 \u2014 IPC Events for Loop Progress', level=3)
    add_code_block(doc, '''// New IPC channels in constants.ts:
TASK_LOOP_ITERATION: 'task:loop:iteration',
TASK_LOOP_GATE_RESULT: 'task:loop:gate-result',
TASK_LOOP_ESCALATION: 'task:loop:escalation',
TASK_LOOP_COMPLETE: 'task:loop:complete',''')

    # Specialist Reviews for Phase 2
    doc.add_heading('Specialist Reviews \u2014 Phase 2', level=2)

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'CRITICAL: Worktree must be reused across task loop iterations. Currently each specialist spawn creates '
        'a fresh worktree, losing all prior work from previous iterations. Pass the SAME worktree path to all '
        'iterations of the same task. Only create a new worktree on the first iteration.'
    )
    doc.add_paragraph(
        'Fix context must be capped at ~1000 tokens. After 5+ iterations, accumulated fix context can exceed '
        'the context window. Keep only the last 2 iteration results. Summarize older iterations into a single '
        'paragraph.'
    )
    doc.add_paragraph(
        'Stuck detection should also check for oscillating errors (A->B->A->B pattern), not just identical errors.'
    )

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'Quality gate commands run as child processes inside the user\'s workspace. Use child_process.execFile '
        '(not exec) to avoid shell injection. Set a per-gate timeout (default 120s). Gate processes must inherit '
        'the workspace\'s PATH and node_modules/.bin. Use process.cwd = worktreePath.'
    )
    doc.add_paragraph(
        'IPC event flooding: Task loop can emit many events per second during gate execution. Throttle renderer '
        'updates to 200ms minimum interval. Batch multiple gate results into a single IPC message.'
    )

    add_specialist_badge(doc, 'UX/UI Specialist', '\U0001F3A8')
    doc.add_paragraph('New component: TaskLoopProgress')
    doc.add_paragraph(
        'Displays inside the existing TaskProgress panel. Shows: iteration counter ("Attempt 2/10"), '
        'model tier badge (haiku=emerald, sonnet=blue, opus=purple), gate results grid (4 columns: Tests, '
        'Types, Lint, Security \u2014 each with pass/fail icon), escalation timeline (vertical line with model '
        'change markers). Animated pulse on the current active gate. Collapsible error output panels.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS task_loop_iterations (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  agent_session_id TEXT NOT NULL REFERENCES agent_sessions(id),
  iteration INTEGER NOT NULL,
  model_used TEXT NOT NULL,
  gates_passed INTEGER NOT NULL DEFAULT 0,
  failed_gates TEXT,              -- JSON array: ["tests","lint"]
  gate_output TEXT,               -- full output for debugging
  error_summary TEXT,             -- truncated error for fix context
  escalated INTEGER DEFAULT 0,
  escalated_from TEXT,
  escalated_to TEXT,
  duration_ms INTEGER,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_task_loop_session ON task_loop_iterations(agent_session_id);
CREATE INDEX idx_task_loop_model ON task_loop_iterations(model_used);''')

    doc.add_heading('Files to Create/Modify', level=2)
    add_styled_table(doc,
        ['Action', 'File'],
        [
            ['CREATE', 'src/main/services/quality-gate.service.ts'],
            ['CREATE', 'src/main/services/task-loop.service.ts'],
            ['MODIFY', 'src/main/services/specialist-pool.service.ts \u2014 integrate task loop'],
            ['MODIFY', 'src/shared/constants.ts \u2014 add loop IPC channels'],
            ['MODIFY', 'src/shared/types.ts \u2014 add TaskLoopState, GateResult, QualityGateReport'],
            ['MODIFY', 'src/main/db/schema.sql \u2014 add task_loop_iterations table'],
            ['CREATE', 'src/main/db/repositories/task-loop.repository.ts'],
            ['MODIFY', 'src/renderer/src/components/agents/TaskProgress.tsx \u2014 show loop state'],
        ])

    doc.add_page_break()

    # =====================================================================
    # PHASE 3
    # =====================================================================
    doc.add_heading('5. Phase 3: Anti-Abandonment Detection', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Detect when a specialist is giving up ("I cannot", "this is beyond my capabilities") and inject '
        'a re-engagement prompt instead of accepting the failure.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Persistence hook (regex)', 'DevTeam', 'hooks/persistence-hook.sh'],
            ['Persistence config', 'DevTeam', '.devteam/persistence-config.yaml'],
        ])

    doc.add_heading('Pattern Categories', level=2)
    doc.add_paragraph('From DevTeam persistence-hook.sh:')

    patterns = [
        ('DIRECT ABANDONMENT', '"I cannot complete", "I\'m unable to", "I give up", "beyond my capabilities"'),
        ('PREMATURE COMPLETION', '"I\'ve done what I can", "I\'ve tried everything", "I\'m out of ideas"'),
        ('DEFLECTION TO USER', '"You should try", "This requires human", "You\'ll need to manually"'),
        ('PERMISSION SEEKING', '"Should I proceed", "Do you want me to", "Would you like me to"'),
        ('LEGITIMATE STOP (whitelist)', '"All tests passing", "Task completed successfully", "Implementation complete"'),
    ]
    for cat, examples in patterns:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{cat}: ')
        run.bold = True
        p.add_run(examples)

    doc.add_heading('What to Build', level=2)
    doc.add_paragraph('Create: src/main/services/abandonment-detector.service.ts')
    add_code_block(doc, '''export interface AbandonmentDetection {
  detected: boolean
  category: 'direct' | 'premature' | 'deflection' | 'permission' | null
  matchedPhrase: string | null
  confidence: number  // 0-1
}

export class AbandonmentDetector {
  detect(text: string): AbandonmentDetection
  getReEngagementPrompt(detection: AbandonmentDetection, taskDescription: string): string
}''')

    # Specialist Reviews for Phase 3
    doc.add_heading('Specialist Reviews \u2014 Phase 3', level=2)

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'CRITICAL (P0): Specialists have stdin: \'ignore\' in their spawn config. You CANNOT write '
        're-engagement prompts to their stdin. For specialists, the approach must be: detect abandonment '
        'in the accumulated text AFTER the process exits, then spawn a NEW process with the re-engagement '
        'context prepended to the task description. This integrates naturally with the Task Loop (Phase 2) \u2014 '
        'abandonment is just another type of iteration failure.'
    )
    doc.add_paragraph(
        'For the generalist (long-lived session with stdin pipe), re-engagement CAN be injected via stdin. '
        'But the format must be: { type: "user", message: { role: "user", content: prompt } } \u2014 NOT the '
        'simplified format shown in the plan.'
    )

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'The regex patterns must be case-insensitive and handle multi-line text. Test for false positives: '
        'e.g., a specialist discussing "what a human should consider" in documentation should NOT trigger '
        'the deflection detector. Add context-awareness: only flag patterns in the LAST 500 characters '
        '(the conclusion) of specialist output, not in the middle of implementation discussion.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    add_code_block(doc, '''-- Phase 3: Track abandonment events
ALTER TABLE agent_sessions ADD COLUMN abandonment_count INTEGER DEFAULT 0;
ALTER TABLE agent_sessions ADD COLUMN abandonment_categories TEXT; -- JSON array

-- Optionally track in task_loop_iterations
ALTER TABLE task_loop_iterations ADD COLUMN abandonment_detected INTEGER DEFAULT 0;
ALTER TABLE task_loop_iterations ADD COLUMN abandonment_category TEXT;''')

    doc.add_heading('Files to Create/Modify', level=2)
    add_styled_table(doc,
        ['Action', 'File'],
        [
            ['CREATE', 'src/main/services/abandonment-detector.service.ts'],
            ['MODIFY', 'src/main/services/specialist-pool.service.ts \u2014 check before marking complete'],
            ['MODIFY', 'src/main/services/generalist.service.ts \u2014 detect in accumulated text'],
            ['MODIFY', 'src/shared/constants.ts \u2014 add ABANDONMENT_DETECTED IPC channel'],
            ['MODIFY', 'src/shared/types.ts \u2014 add AbandonmentDetection interface'],
            ['MODIFY', 'src/renderer/src/components/chat/MessageBubble.tsx \u2014 show re-engagement indicator'],
        ])

    doc.add_page_break()

    # =====================================================================
    # PHASE 4
    # =====================================================================
    doc.add_heading('6. Phase 4: File-Based Agent Communication Chain', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Instead of passing all context through the LLM context window, have each specialist write structured '
        'output to a shared task directory. Subsequent specialists read prior outputs explicitly. This makes '
        'workflows auditable, resumable, and debuggable.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['File chain (.full-stack-feature/)', 'wshobson/agents', 'plugins/full-stack-orchestration/commands/full-stack-feature.md'],
            ['State.json pattern', 'wshobson/agents', 'Same file (state management section)'],
            ['PROJECT_STATUS.md', 'Multi-Agent Squad', 'Root directory pattern'],
        ])

    doc.add_heading('Artifact Directory Structure', level=2)
    add_code_block(doc, '''{workspace}/.agentstudio/
  {conversation-id}/
    state.json                    # Overall execution state
    tasks/
      {task-id}/
        input.md                  # Task description + prior context
        output.md                 # Specialist's structured output
        gate-results.json         # Quality gate results (Phase 2)
        iterations.json           # Task loop history (Phase 2)''')

    doc.add_heading('What to Build', level=2)
    doc.add_paragraph('Create: src/main/services/task-artifact.service.ts')
    add_code_block(doc, '''export interface TaskArtifactState {
  planId: string
  conversationId: string
  status: 'in_progress' | 'complete' | 'failed' | 'paused'
  currentTaskIndex: number
  completedTasks: string[]
  failedTasks: string[]
  artifacts: Record<string, string>  // taskId -> output file path
  startedAt: string
  lastUpdated: string
}

export class TaskArtifactService {
  async initialize(workspacePath: string, conversationId: string, plan: TaskPlan): Promise<void>
  async writeTaskInput(taskId: string, content: string): Promise<string>
  async writeTaskOutput(taskId: string, content: string): Promise<string>
  async readPriorOutputs(taskId: string, dependencies: string[]): Promise<string>
  async updateState(update: Partial<TaskArtifactState>): Promise<void>
  async readState(workspacePath: string, conversationId: string): Promise<TaskArtifactState | null>
  async cleanup(workspacePath: string, conversationId: string): Promise<void>
}''')

    # Specialist Reviews
    doc.add_heading('Specialist Reviews \u2014 Phase 4', level=2)

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'File operations must use path.join() consistently. The .agentstudio directory must be created with '
        'fs.mkdir recursive:true. File writes should be atomic (write to temp file, then rename) to prevent '
        'corruption if the app crashes mid-write. Add .agentstudio/ to the workspace\'s .gitignore automatically.'
    )

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'The enriched prompt template for specialists should include explicit output formatting instructions. '
        'Each specialist should write output.md with a consistent structure: ## Summary, ## Files Modified, '
        '## Key Decisions, ## API Signatures, ## Test Coverage. This makes downstream parsing reliable.'
    )

    doc.add_page_break()

    # =====================================================================
    # PHASE 5
    # =====================================================================
    doc.add_heading('7. Phase 5: Cost Tracking Dashboard', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Track token usage and estimated cost per model, per agent, per conversation. Show users how much '
        'they\'re spending and how much model routing (Phase 1) is saving them.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Cost tracking tables', 'DevTeam', 'scripts/schema.sql (sessions table)'],
            ['Cost tracking script', 'DevTeam', 'scripts/cost-tracking.sh'],
        ])

    doc.add_heading('What to Build', level=2)
    doc.add_paragraph('Create: src/main/services/cost-calculator.service.ts')
    add_code_block(doc, '''export interface CostBreakdown {
  totalCostCents: number
  byModel: Record<string, { inputTokens: number, outputTokens: number, costCents: number }>
  byAgent: Record<string, { inputTokens: number, outputTokens: number, costCents: number }>
  savingsVsAllOpus: number  // estimated savings from model routing
}

export class CostCalculator {
  calculateSessionCost(tokens: { input: number, output: number }, model: string): number
  getConversationCost(conversationId: string): Promise<CostBreakdown>
  getWorkspaceCost(workspaceId: string, since?: Date): Promise<CostBreakdown>
  calculateSavings(breakdown: CostBreakdown): number
}''')

    # Specialist Reviews
    doc.add_heading('Specialist Reviews \u2014 Phase 5', level=2)

    add_specialist_badge(doc, 'UX/UI Specialist', '\U0001F3A8')
    doc.add_paragraph('New component: CostDashboard')
    doc.add_paragraph(
        'Full-page dashboard in Workspace Settings with four sections: '
        '(1) Total spend card with large dollar amount and period selector (7d/30d/all), '
        '(2) Model spend breakdown as horizontal stacked bars (haiku=emerald, sonnet=blue, opus=purple) with '
        'token counts and cost labels, '
        '(3) Savings highlight card showing "You saved $X.XX vs all-Opus" with percentage badge, '
        '(4) Per-conversation cost table with columns: Conversation, Messages, Tokens, Cost, Model Mix. '
        'Budget alert banner at top when >80% of budget used (amber-500 background).'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    doc.add_paragraph('Repository: src/main/db/repositories/cost.repository.ts')
    add_code_block(doc, '''export class CostRepository {
  getConversationCost(conversationId: string): CostBreakdown {
    // SQL: JOIN agent_sessions with model_pricing
    // GROUP BY model_used, agent_id
    // SUM tokens_input * input_per_1m / 1000000
  }

  getWorkspaceCost(workspaceId: string, since?: string): CostBreakdown {
    // Similar with workspace_id filter and date range
  }

  getDailyCostTrend(workspaceId: string, days: number): DailyCost[] {
    // GROUP BY date(created_at) for sparkline chart
  }
}''')

    doc.add_heading('Files to Create/Modify', level=2)
    add_styled_table(doc,
        ['Action', 'File'],
        [
            ['CREATE', 'src/main/services/cost-calculator.service.ts'],
            ['CREATE', 'src/main/ipc/cost.ipc.ts'],
            ['CREATE', 'src/renderer/src/components/settings/CostDashboard.tsx'],
            ['CREATE', 'src/renderer/src/store/cost.store.ts'],
            ['CREATE', 'src/main/db/repositories/cost.repository.ts'],
            ['MODIFY', 'src/shared/constants.ts \u2014 add cost IPC channels'],
            ['MODIFY', 'src/shared/types.ts \u2014 add CostBreakdown, DailyCost types'],
            ['MODIFY', 'src/renderer/src/components/agents/TaskProgress.tsx \u2014 per-task cost'],
        ])

    doc.add_page_break()

    # =====================================================================
    # PHASE 6
    # =====================================================================
    doc.add_heading('8. Phase 6: Human Checkpoint UI', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Add explicit approval gates before critical actions: merging worktrees, executing multi-task plans, '
        'deploying. The user sees a summary of what will happen and must approve.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Phase checkpoints', 'wshobson/agents', 'plugins/full-stack-orchestration/commands/full-stack-feature.md'],
            ['CRITICAL DECISION pattern', 'Multi-Agent Squad', '.claude/hooks/enterprise-workflow.toml'],
            ['Conductor approval gates', 'wshobson/agents', 'plugins/conductor/commands/implement.md'],
        ])

    doc.add_heading('What to Build', level=2)
    add_code_block(doc, '''export interface Checkpoint {
  id: string
  type: 'phase_gate' | 'merge_approval' | 'deployment' | 'destructive_action'
  title: string
  summary: string
  details: {
    what: string          // What will happen
    why: string           // Why this checkpoint exists
    risk: string          // What could go wrong
    changedFiles?: string[]
    testResults?: string
  }
  status: 'pending' | 'approved' | 'rejected'
  createdAt: string
}

export class CheckpointService extends EventEmitter {
  async requestApproval(checkpoint: Omit<Checkpoint, 'id' | 'status' | 'createdAt'>): Promise<boolean>
  async resolve(checkpointId: string, approved: boolean): Promise<void>
}''')

    doc.add_paragraph(
        'When requestApproval is called, it: (1) Creates the checkpoint record, (2) Emits an IPC event to '
        'the renderer, (3) Returns a Promise that resolves when the user approves/rejects, (4) The specialist '
        'pool PAUSES at this point.'
    )

    # Specialist Reviews
    doc.add_heading('Specialist Reviews \u2014 Phase 6', level=2)

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'The Promise-based blocking pattern (requestApproval returns a Promise that resolves on user action) '
        'is correct for IPC. Implement with a Map<checkpointId, { resolve, reject }> in the main process. '
        'When the renderer calls resolve(checkpointId, approved), look up and resolve the pending Promise. '
        'Add a timeout (default: 30 minutes) to prevent orphaned Promises if the user closes the window.'
    )

    add_specialist_badge(doc, 'UX/UI Specialist', '\U0001F3A8')
    doc.add_paragraph('New component: CheckpointModal')
    doc.add_paragraph(
        'Full-screen modal overlay with three sections: WHAT (plain language description of the action), '
        'WHY (reason this checkpoint exists), RISK (what could go wrong, styled with amber-100 background). '
        'Below: collapsible file list showing all changed files with diff stats (+/- lines). '
        'Two buttons: "Approve" (emerald-600, left-aligned) and "Reject" (red-600, right-aligned) with '
        'confirmation for reject. When rejected, show a text input for feedback that gets passed back '
        'to the specialist as context for the next attempt.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS checkpoints (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  conversation_id TEXT NOT NULL REFERENCES conversations(id),
  type TEXT NOT NULL,
  title TEXT NOT NULL,
  summary TEXT,
  details_json TEXT,              -- JSON: { what, why, risk, changedFiles, testResults }
  status TEXT NOT NULL DEFAULT 'pending',
  resolved_at TEXT,
  user_feedback TEXT,             -- Feedback on rejection
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_checkpoints_conversation ON checkpoints(conversation_id);
CREATE INDEX idx_checkpoints_status ON checkpoints(status);''')

    doc.add_page_break()

    # =====================================================================
    # PHASE 7
    # =====================================================================
    doc.add_heading('9. Phase 7: Progressive Skill Loading', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Stop dumping all specialist knowledge into every prompt. Load skills in tiers: metadata always available, '
        'instructions on activation, code examples on demand. Reduces token usage and improves focus.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Progressive disclosure (146 skills)', 'wshobson/agents', 'docs/agent-skills.md'],
            ['Skill activation matching', 'wshobson/agents', 'Built-in description matching'],
        ])

    doc.add_heading('Tier Structure', level=2)
    add_code_block(doc, '''// Tier 1 (Always loaded, ~50 tokens): Metadata
//   name, description, activation keywords
//
// Tier 2 (On activation, ~500 tokens): Core instructions
//   Key patterns, decision rules, approach guidance
//
// Tier 3 (On demand, full content): Resources
//   Code examples, templates, reference implementations

export interface SkillTier {
  tier1: { name: string; description: string; keywords: string[] }
  tier2: string | null     // First section of the skill
  tier3: string            // Full content
}''')

    doc.add_heading('Smart Matching Strategy', level=2)
    doc.add_paragraph(
        '1. First pass: Match against Tier 1 keywords (fast, no LLM call). '
        '2. If multiple matches, use LLM with Tier 2 content to pick best match. '
        '3. Inject Tier 2 by default. Only inject Tier 3 if task specifically needs code examples.'
    )

    doc.add_heading('Files to Create/Modify', level=2)
    add_styled_table(doc,
        ['Action', 'File'],
        [
            ['MODIFY', 'src/main/services/skill.service.ts \u2014 add tier parsing, smart loading'],
            ['MODIFY', 'src/main/services/orchestrator.service.ts \u2014 tiered matching'],
            ['MODIFY', 'src/main/db/schema.sql \u2014 add tier columns to skills table'],
            ['MODIFY', 'src/main/db/repositories/skill.repository.ts \u2014 persist tiers'],
            ['MODIFY', 'src/shared/types.ts \u2014 add SkillTier interface'],
        ])

    doc.add_page_break()

    # =====================================================================
    # PHASE 8
    # =====================================================================
    doc.add_heading('10. Phase 8: Scope Enforcement Layer', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Prevent specialists from modifying files outside their assigned task scope. Add file-pattern '
        'restrictions on top of the existing worktree isolation.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Scope validator (6-layer VETO)', 'DevTeam', 'agents/orchestration/scope-validator.md'],
            ['Scope check hook', 'DevTeam', 'hooks/scope-check.sh'],
        ])

    doc.add_heading('What to Build', level=2)
    add_code_block(doc, '''export interface TaskScope {
  allowedFiles?: string[]        // Exact paths
  allowedPatterns?: string[]     // Globs
  forbiddenDirectories?: string[]
  forbiddenFiles?: string[]
  maxFilesChanged?: number
}

export class ScopeValidator {
  validate(changedFiles: string[], scope: TaskScope): ScopeValidationResult
  async getChangedFiles(worktreePath: string): Promise<string[]>
  async validatePostExecution(worktreePath: string, scope: TaskScope): Promise<ScopeValidationResult>
}

// Validation priority:
//   1. Check forbidden files FIRST (highest priority)
//   2. Check forbidden directories
//   3. Check explicitly allowed files
//   4. Check glob patterns
//   5. NOT explicitly allowed = FAIL''')

    doc.add_heading('Specialist Reviews \u2014 Phase 8', level=2)

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'CRITICAL: Scope validation MUST happen BEFORE worktree merge, not after. Currently the specialist pool '
        'auto-merges on completion. Validation sequence must be: specialist completes -> scope validation -> '
        'quality gates -> checkpoint approval -> THEN merge. If scope violation detected, revert violating files '
        '(git checkout -- <files>) and re-run the specialist with scope instructions reinforced.'
    )

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'Use simple-git library (already a dependency) for git diff --name-only and git checkout operations. '
        'Glob matching should use minimatch or picomatch for consistent pattern behavior.'
    )

    doc.add_page_break()

    # =====================================================================
    # PHASE 9
    # =====================================================================
    doc.add_heading('11. Phase 9: Declarative Hooks System', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Let users define automation hooks in a config file (YAML) instead of hardcoding behavior. '
        'Hooks trigger on events and run shell commands or inject prompts.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['TOML declarative hooks', 'Multi-Agent Squad', '.claude/hooks/enterprise-workflow.toml'],
            ['Dynamic hook generation', 'Multi-Agent Squad', 'scripts/generate-hooks.py'],
            ['JSON hook config', 'DevTeam', 'hooks/hooks.json'],
        ])

    doc.add_heading('Hook Configuration Format', level=2)
    add_code_block(doc, '''# .agentstudio/hooks.yaml
hooks:
  - event: specialist_complete
    name: "Auto-format code"
    command: "npx prettier --write ."
    blocking: false

  - event: gate_failed
    name: "Notify on test failure"
    command: "echo 'Tests failed' | notify-send"
    blocking: false

  - event: pre_merge
    name: "Final lint check"
    command: "npm run lint"
    blocking: true''')

    doc.add_heading('Hook Events', level=2)
    events = [
        'specialist_start / specialist_complete / specialist_failed',
        'gate_passed / gate_failed',
        'escalation',
        'pre_merge / post_merge',
        'plan_created',
        'checkpoint_approved / checkpoint_rejected',
        'abandonment_detected',
        'task_loop_complete',
    ]
    for e in events:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('Specialist Reviews \u2014 Phase 9', level=2)

    add_specialist_badge(doc, 'Electron Architect', '\u26A1')
    doc.add_paragraph(
        'SECURITY: Hook commands run as shell commands in the user\'s workspace. Use child_process.execFile '
        'with explicit argument splitting, NOT child_process.exec (shell injection risk). Validate the YAML '
        'schema before loading. Limit environment variable interpolation to a whitelist of safe variables. '
        'Never pass user-controlled data directly into command strings.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS hook_executions (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  workspace_id TEXT NOT NULL REFERENCES workspaces(id),
  hook_name TEXT NOT NULL,
  event TEXT NOT NULL,
  command TEXT NOT NULL,
  exit_code INTEGER,
  output TEXT,
  duration_ms INTEGER,
  blocking INTEGER DEFAULT 0,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_hook_exec_workspace ON hook_executions(workspace_id);
CREATE INDEX idx_hook_exec_event ON hook_executions(event);''')

    doc.add_page_break()

    # =====================================================================
    # PHASE 10
    # =====================================================================
    doc.add_heading('12. Phase 10: Deep Agent Personas & Bug Council', level=1)

    doc.add_heading('Goal', level=2)
    doc.add_paragraph(
        'Enrich specialist definitions with deep expertise (war stories, red flags, philosophy). '
        'Add a Bug Council pattern where 5 diagnostic agents collaborate when a specialist fails at opus.'
    )

    doc.add_heading('Reference Code', level=2)
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        [
            ['Deep agent personas (100+ lines)', 'Multi-Agent Squad', '.claude/agents/*/*.md'],
            ['Bug Council (5 analysts)', 'DevTeam', 'agents/diagnosis/*.md'],
            ['Bug Council orchestrator', 'DevTeam', 'agents/orchestration/bug-council-orchestrator.md'],
        ])

    doc.add_heading('Deep Persona Template', level=2)
    add_code_block(doc, '''## Role
You are a {role} with {years}+ years of experience building production systems.

## Core Expertise
{expertise_bullets}

## Lessons from Production
{war_stories}

## Red Flags You Catch
{red_flags}

## Your Approach
{code_patterns}

## Quality Commitments
- I never ship code without tests
- I always check for security implications
- I document non-obvious decisions
- I keep changes focused and minimal''')

    doc.add_heading('Bug Council \u2014 5 Diagnostic Perspectives', level=2)
    council = [
        ('Root Cause Analyst', 'Error analysis, hypothesis generation, stack trace dissection'),
        ('Code Archaeologist', 'Git history, regression detection, blame analysis'),
        ('Pattern Matcher', 'Similar bugs in codebase, anti-pattern identification'),
        ('Systems Thinker', 'Dependency analysis, architectural issues, cascade effects'),
        ('Adversarial Tester', 'Edge cases, security vulnerabilities, concurrency bugs'),
    ]
    for i, (name, desc) in enumerate(council, 1):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{i}. {name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph(
        'Activation triggers: 3+ consecutive opus failures, explicit user request, or stuck loop detection. '
        'Implementation: Spawn 5 parallel claude -p calls (haiku for cost efficiency), each with a different '
        'diagnostic persona. Collect outputs, then make a final synthesis call (sonnet) to combine perspectives '
        'into an actionable solution.'
    )

    # Specialist Reviews
    doc.add_heading('Specialist Reviews \u2014 Phase 10', level=2)

    add_specialist_badge(doc, 'Agentic Architect', '\U0001F916')
    doc.add_paragraph(
        'Bug Council should use haiku for the 5 parallel diagnostic calls (text analysis, not code generation) '
        'and sonnet for the synthesis. This keeps cost under $0.05 per council session. The synthesis prompt '
        'should explicitly ask for: (1) most likely root cause, (2) recommended fix approach, (3) files to examine, '
        '(4) test to write to verify the fix.'
    )

    add_specialist_badge(doc, 'UX/UI Specialist', '\U0001F3A8')
    doc.add_paragraph('New component: BugCouncilPanel')
    doc.add_paragraph(
        'Tabbed interface with 5 tabs (one per analyst) plus a "Synthesis" tab. Each tab shows the analyst\'s '
        'findings with syntax-highlighted code snippets where relevant. The Synthesis tab has a highlighted '
        'recommendation box (blue-50 background) with the recommended approach. Two action buttons: '
        '"Retry with Council Guidance" (primary) and "Escalate to Human" (secondary). '
        'Shows a confidence meter (0-100%) based on agreement between analysts.'
    )

    add_specialist_badge(doc, 'DB Architect', '\U0001F5C4')
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS bug_council_sessions (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  agent_session_id TEXT NOT NULL REFERENCES agent_sessions(id),
  trigger_reason TEXT NOT NULL,
  task_description TEXT,
  failure_count INTEGER,
  synthesized_solution TEXT,
  recommended_approach TEXT,
  outcome TEXT,                   -- 'resolved' | 'escalated_to_human' | 'failed'
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS bug_council_analyses (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  council_session_id TEXT NOT NULL REFERENCES bug_council_sessions(id),
  perspective TEXT NOT NULL,
  findings TEXT NOT NULL,
  confidence REAL,
  model_used TEXT,
  tokens_used INTEGER,
  duration_ms INTEGER,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_bug_council_session ON bug_council_sessions(agent_session_id);
CREATE INDEX idx_bug_analyses_council ON bug_council_analyses(council_session_id);''')

    doc.add_page_break()

    # =====================================================================
    # 13. DATABASE MIGRATION GUIDE
    # =====================================================================
    doc.add_heading('13. Database Migration Guide (Complete SQL)', level=1)

    doc.add_paragraph(
        'All database changes across all 10 phases, organized into a single migration script. '
        'Apply as a single transaction in src/main/db/index.ts.'
    )

    doc.add_heading('Migration Strategy', level=2)
    doc.add_paragraph(
        'Add a schema_version table to track applied migrations. Each migration has a version number '
        'and is applied only once. Wrap all operations in try/catch to maintain the existing AgentStudio pattern.'
    )

    add_code_block(doc, '''-- Migration infrastructure
CREATE TABLE IF NOT EXISTS schema_version (
  version INTEGER PRIMARY KEY,
  description TEXT NOT NULL,
  applied_at TEXT NOT NULL DEFAULT (datetime('now'))
);''')

    doc.add_heading('Migration v1: Complexity & Cost (Phases 1, 5)', level=3)
    add_code_block(doc, '''-- agent_sessions extensions
ALTER TABLE agent_sessions ADD COLUMN complexity_score INTEGER;
ALTER TABLE agent_sessions ADD COLUMN complexity_tier TEXT;
ALTER TABLE agent_sessions ADD COLUMN model_used TEXT;
ALTER TABLE agent_sessions ADD COLUMN tokens_input INTEGER DEFAULT 0;
ALTER TABLE agent_sessions ADD COLUMN tokens_output INTEGER DEFAULT 0;

-- Model pricing
CREATE TABLE IF NOT EXISTS model_pricing (
  model_id TEXT PRIMARY KEY,
  input_per_1m_cents INTEGER NOT NULL,
  output_per_1m_cents INTEGER NOT NULL,
  effective_date TEXT NOT NULL DEFAULT (datetime('now'))
);

INSERT OR IGNORE INTO model_pricing VALUES
  ('claude-haiku-4-20250414', 80, 400, datetime('now')),
  ('claude-sonnet-4-20250514', 300, 1500, datetime('now')),
  ('claude-opus-4-20250514', 1500, 7500, datetime('now'));''')

    doc.add_heading('Migration v2: Task Loop (Phase 2)', level=3)
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS task_loop_iterations (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  agent_session_id TEXT NOT NULL REFERENCES agent_sessions(id),
  iteration INTEGER NOT NULL,
  model_used TEXT NOT NULL,
  gates_passed INTEGER NOT NULL DEFAULT 0,
  failed_gates TEXT,
  gate_output TEXT,
  error_summary TEXT,
  escalated INTEGER DEFAULT 0,
  escalated_from TEXT,
  escalated_to TEXT,
  abandonment_detected INTEGER DEFAULT 0,
  abandonment_category TEXT,
  duration_ms INTEGER,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_task_loop_session ON task_loop_iterations(agent_session_id);
CREATE INDEX idx_task_loop_model ON task_loop_iterations(model_used);''')

    doc.add_heading('Migration v3: Anti-Abandonment (Phase 3)', level=3)
    add_code_block(doc, '''ALTER TABLE agent_sessions ADD COLUMN abandonment_count INTEGER DEFAULT 0;
ALTER TABLE agent_sessions ADD COLUMN abandonment_categories TEXT;''')

    doc.add_heading('Migration v4: Checkpoints (Phase 6)', level=3)
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS checkpoints (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  conversation_id TEXT NOT NULL REFERENCES conversations(id),
  type TEXT NOT NULL,
  title TEXT NOT NULL,
  summary TEXT,
  details_json TEXT,
  status TEXT NOT NULL DEFAULT 'pending',
  resolved_at TEXT,
  user_feedback TEXT,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_checkpoints_conversation ON checkpoints(conversation_id);
CREATE INDEX idx_checkpoints_status ON checkpoints(status);''')

    doc.add_heading('Migration v5: Skills Tier (Phase 7)', level=3)
    add_code_block(doc, '''ALTER TABLE skills ADD COLUMN tier1_summary TEXT;
ALTER TABLE skills ADD COLUMN tier2_instructions TEXT;''')

    doc.add_heading('Migration v6: Hooks (Phase 9)', level=3)
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS hook_executions (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  workspace_id TEXT NOT NULL REFERENCES workspaces(id),
  hook_name TEXT NOT NULL,
  event TEXT NOT NULL,
  command TEXT NOT NULL,
  exit_code INTEGER,
  output TEXT,
  duration_ms INTEGER,
  blocking INTEGER DEFAULT 0,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_hook_exec_workspace ON hook_executions(workspace_id);
CREATE INDEX idx_hook_exec_event ON hook_executions(event);''')

    doc.add_heading('Migration v7: Bug Council (Phase 10)', level=3)
    add_code_block(doc, '''CREATE TABLE IF NOT EXISTS bug_council_sessions (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  agent_session_id TEXT NOT NULL REFERENCES agent_sessions(id),
  trigger_reason TEXT NOT NULL,
  task_description TEXT,
  failure_count INTEGER,
  synthesized_solution TEXT,
  recommended_approach TEXT,
  outcome TEXT,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE TABLE IF NOT EXISTS bug_council_analyses (
  id TEXT PRIMARY KEY DEFAULT (hex(randomblob(16))),
  council_session_id TEXT NOT NULL REFERENCES bug_council_sessions(id),
  perspective TEXT NOT NULL,
  findings TEXT NOT NULL,
  confidence REAL,
  model_used TEXT,
  tokens_used INTEGER,
  duration_ms INTEGER,
  created_at TEXT NOT NULL DEFAULT (datetime('now'))
);

CREATE INDEX idx_bug_council_session ON bug_council_sessions(agent_session_id);
CREATE INDEX idx_bug_analyses_council ON bug_council_analyses(council_session_id);''')

    doc.add_page_break()

    # =====================================================================
    # 14. UX/UI COMPONENT SPECIFICATIONS
    # =====================================================================
    doc.add_heading('14. UX/UI Component Specifications', level=1)

    doc.add_paragraph(
        'All new UI components designed by the UX/UI Specialist. Each includes purpose, location, '
        'layout description, Tailwind CSS 4 classes, and accessibility requirements.'
    )

    components = [
        {
            'name': 'CostPreferenceSelector',
            'phase': 'Phase 1',
            'location': 'Workspace Settings > General',
            'description': (
                '3-segment toggle: Economy (emerald-500) / Balanced (blue-500) / Power (purple-500). '
                'Each segment shows a label and subtitle. Economy: "~40% savings". Balanced: "Recommended". '
                'Power: "Max capability". Selected segment has filled background, others have ghost style. '
                'Keyboard accessible with arrow keys. Stores in workspace settings_json.'
            ),
        },
        {
            'name': 'TaskLoopProgress',
            'phase': 'Phase 2',
            'location': 'TaskProgress panel (existing)',
            'description': (
                'Iteration counter ("Attempt 2/10"), model tier badge with color coding, '
                'gate results grid (4 columns: Tests, Types, Lint, Security with pass/fail icons), '
                'escalation timeline (vertical line with model change markers). '
                'Animated pulse on active gate. Collapsible error output panels per gate.'
            ),
        },
        {
            'name': 'AbandonmentIndicator',
            'phase': 'Phase 3',
            'location': 'Chat message stream',
            'description': (
                'Inline notification in the chat stream: amber-50 background with amber-600 border-left. '
                'Icon: lightning bolt. Text: "Agent attempted to give up - re-engaging with alternative approach". '
                'Collapsible detail showing matched pattern category and re-engagement prompt sent.'
            ),
        },
        {
            'name': 'TaskArtifactBrowser',
            'phase': 'Phase 4',
            'location': 'Conversation detail sidebar',
            'description': (
                'Tree view of .agentstudio/{conversation-id}/ directory. Shows state.json status badge '
                '(in_progress/complete/failed), task list with completion checkmarks, clickable output.md '
                'files that open in a markdown preview panel. "Resume" button for paused workflows.'
            ),
        },
        {
            'name': 'CostDashboard',
            'phase': 'Phase 5',
            'location': 'Workspace Settings > Cost',
            'description': (
                'Full-page dashboard: (1) Total spend card with period selector, '
                '(2) Model spend horizontal stacked bars, (3) Savings highlight card, '
                '(4) Per-conversation cost table. Budget alert banner when >80% used.'
            ),
        },
        {
            'name': 'CostBadge',
            'phase': 'Phase 5',
            'location': 'Conversation header',
            'description': (
                'Small inline badge showing current conversation cost: "$0.12". '
                'Color-coded: gray-500 (<$0.50), amber-500 ($0.50-$2.00), red-500 (>$2.00). '
                'Tooltip shows breakdown by model on hover.'
            ),
        },
        {
            'name': 'CheckpointModal',
            'phase': 'Phase 6',
            'location': 'Global modal overlay',
            'description': (
                'Three sections: WHAT (action description), WHY (checkpoint reason), RISK (amber-100 background). '
                'Collapsible changed files list with diff stats. Approve (emerald-600) and Reject (red-600) buttons. '
                'Reject shows text input for feedback. Blocks execution until resolved.'
            ),
        },
        {
            'name': 'SkillTierIndicator',
            'phase': 'Phase 7',
            'location': 'Agent detail panel',
            'description': (
                'Shows which skills are loaded for the current specialist and at which tier. '
                'Tier 1 = gray dot (metadata), Tier 2 = blue dot (instructions), Tier 3 = purple dot (full). '
                'Expandable to show token count per loaded skill.'
            ),
        },
        {
            'name': 'ScopeViolationAlert',
            'phase': 'Phase 8',
            'location': 'Task progress panel',
            'description': (
                'Red-50 background alert showing scope violations: file path, reason (forbidden/not_allowed), '
                'and action taken (reverted). Collapsible to show full validation result.'
            ),
        },
        {
            'name': 'HookManager',
            'phase': 'Phase 9',
            'location': 'Workspace Settings > Hooks',
            'description': (
                'List of configured hooks with event type, command, blocking status. '
                'Execution history per hook (last run, status, output). '
                '"Generate Defaults" button based on detected project type. '
                'Inline YAML editor for advanced users.'
            ),
        },
        {
            'name': 'BugCouncilPanel',
            'phase': 'Phase 10',
            'location': 'Task detail panel (when activated)',
            'description': (
                'Tabbed interface: 5 analyst tabs + Synthesis tab. Each shows findings with '
                'syntax-highlighted code. Synthesis tab has blue-50 recommendation box. '
                'Confidence meter (0-100%). Action buttons: "Retry with Guidance" / "Escalate to Human".'
            ),
        },
    ]

    for comp in components:
        doc.add_heading(f'{comp["name"]} ({comp["phase"]})', level=2)
        p = doc.add_paragraph()
        run = p.add_run('Location: ')
        run.bold = True
        p.add_run(comp['location'])
        doc.add_paragraph(comp['description'])

    doc.add_page_break()

    # =====================================================================
    # 15. IMPLEMENTATION ORDER & SPRINT PLAN
    # =====================================================================
    doc.add_heading('15. Implementation Order & Sprint Plan', level=1)

    doc.add_heading('Dependency Graph', level=2)
    add_code_block(doc, '''Phase 1: Complexity Scoring & Model Routing    <- NO dependencies, start here
    |
Phase 2: Task Loop with Quality Gates          <- depends on Phase 1
    |
Phase 3: Anti-Abandonment Detection            <- depends on Phase 2
    |
Phase 5: Cost Tracking Dashboard               <- depends on Phase 1

Phase 4: File-Based Agent Communication         <- independent, parallel
Phase 6: Human Checkpoint UI                    <- independent, parallel
Phase 7: Progressive Skill Loading              <- independent, parallel

Phase 8: Scope Enforcement                      <- depends on Phase 2
Phase 9: Declarative Hooks                      <- depends on Phases 2, 3, 6
Phase 10: Bug Council                           <- depends on Phase 2''')

    doc.add_heading('Sprint Plan', level=2)
    add_styled_table(doc,
        ['Sprint', 'Phases', 'Focus', 'Key Deliverables'],
        [
            ['Sprint 1', 'Phase 1 + 5', 'Model routing + cost visibility',
             'ComplexityScorerService, CostCalculator, CostDashboard, CostPreferenceSelector, model_pricing table'],
            ['Sprint 2', 'Phase 2 + 3', 'Task loop + anti-abandonment',
             'QualityGateService, TaskLoopService, AbandonmentDetector, TaskLoopProgress, task_loop_iterations table'],
            ['Sprint 3', 'Phase 4 + 6', 'Artifacts + checkpoints',
             'TaskArtifactService, CheckpointService, CheckpointModal, TaskArtifactBrowser, checkpoints table'],
            ['Sprint 4', 'Phase 7 + 8', 'Skills + scope enforcement',
             'SkillTier parsing, ScopeValidator, scope auto-generation, SkillTierIndicator'],
            ['Sprint 5', 'Phase 9 + 10', 'Hooks + Bug Council',
             'HookEngine, BugCouncilService, HookManager, BugCouncilPanel, hook_executions & bug_council tables'],
        ])

    doc.add_heading('Prerequisites (Before Sprint 1)', level=2)
    prereqs = [
        'Implement schema_version migration system in src/main/db/index.ts',
        'Create ProcessPoolService for global concurrency limiting (max 4 concurrent Claude CLI processes)',
        'Fix specialist token tracking (returns 0) in specialist-pool.service.ts',
        'Split token_usage column from TEXT/JSON to tokens_input + tokens_output INTEGER columns',
        'Verify --model flag works correctly with Claude CLI in print mode (-p)',
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # APPENDIX A: FILE INDEX
    # =====================================================================
    doc.add_heading('Appendix A: File Creation/Modification Index', level=1)

    doc.add_heading('New Files to Create', level=2)
    new_files = [
        ('src/main/services/complexity-scorer.service.ts', '1'),
        ('src/main/services/quality-gate.service.ts', '2'),
        ('src/main/services/task-loop.service.ts', '2'),
        ('src/main/services/abandonment-detector.service.ts', '3'),
        ('src/main/services/task-artifact.service.ts', '4'),
        ('src/main/services/cost-calculator.service.ts', '5'),
        ('src/main/services/checkpoint.service.ts', '6'),
        ('src/main/services/scope-validator.service.ts', '8'),
        ('src/main/services/hook-engine.service.ts', '9'),
        ('src/main/services/bug-council.service.ts', '10'),
        ('src/main/services/process-pool.service.ts', 'Pre'),
        ('src/main/ipc/cost.ipc.ts', '5'),
        ('src/main/db/repositories/task-loop.repository.ts', '2'),
        ('src/main/db/repositories/cost.repository.ts', '5'),
        ('src/main/db/repositories/checkpoint.repository.ts', '6'),
        ('src/main/db/repositories/hook.repository.ts', '9'),
        ('src/renderer/src/components/settings/CostPreferenceSelector.tsx', '1'),
        ('src/renderer/src/components/settings/CostDashboard.tsx', '5'),
        ('src/renderer/src/components/settings/HookManager.tsx', '9'),
        ('src/renderer/src/components/chat/CheckpointModal.tsx', '6'),
        ('src/renderer/src/components/agents/BugCouncilPanel.tsx', '10'),
        ('src/renderer/src/store/cost.store.ts', '5'),
        ('src/renderer/src/store/checkpoint.store.ts', '6'),
    ]
    add_styled_table(doc,
        ['File Path', 'Phase'],
        new_files)

    doc.add_heading('Existing Files to Modify', level=2)
    mod_files = [
        ('src/shared/types.ts', '1, 2, 3, 4, 5, 6, 7, 8, 9, 10'),
        ('src/shared/constants.ts', '1, 2, 3, 5, 6, 9, 10'),
        ('src/main/services/specialist-pool.service.ts', '1, 2, 3, 4, 6, 8, 9'),
        ('src/main/services/orchestrator.service.ts', '1, 4, 7, 8'),
        ('src/main/services/generalist.service.ts', '3'),
        ('src/main/services/system-prompts.ts', '8, 10'),
        ('src/main/services/skill.service.ts', '7'),
        ('src/main/db/schema.sql', '1, 2, 3, 5, 6, 7, 9, 10'),
        ('src/main/db/index.ts', 'Pre (migration system)'),
        ('src/main/db/repositories/agent-session.repository.ts', '1, 5'),
        ('src/main/db/repositories/skill.repository.ts', '7'),
        ('src/main/ipc/chat.ipc.ts', '4, 6'),
        ('src/preload/index.ts', '5, 6, 9, 10'),
        ('src/renderer/src/components/agents/TaskProgress.tsx', '2, 5'),
        ('src/renderer/src/components/chat/MessageBubble.tsx', '3'),
    ]
    add_styled_table(doc,
        ['File Path', 'Phases'],
        mod_files)

    doc.add_page_break()

    # =====================================================================
    # APPENDIX B: IPC CHANNEL REGISTRY
    # =====================================================================
    doc.add_heading('Appendix B: New IPC Channel Registry', level=1)

    channels = [
        ('task:loop:iteration', '2', 'Task loop iteration progress'),
        ('task:loop:gate-result', '2', 'Individual gate pass/fail result'),
        ('task:loop:escalation', '2', 'Model escalation event'),
        ('task:loop:complete', '2', 'Task loop finished'),
        ('abandonment:detected', '3', 'Abandonment pattern matched'),
        ('cost:get:conversation', '5', 'Get conversation cost breakdown'),
        ('cost:get:workspace', '5', 'Get workspace cost breakdown'),
        ('cost:get:summary', '5', 'Get cost summary for header badge'),
        ('checkpoint:request', '6', 'Request user approval'),
        ('checkpoint:resolve', '6', 'User approves/rejects checkpoint'),
        ('checkpoint:list', '6', 'List pending checkpoints'),
        ('hook:executed', '9', 'Hook execution completed'),
        ('hook:list', '9', 'List configured hooks'),
        ('hook:generate-defaults', '9', 'Generate default hooks'),
        ('bug-council:activated', '10', 'Bug Council session started'),
        ('bug-council:analysis', '10', 'Individual analyst result'),
        ('bug-council:synthesis', '10', 'Synthesized solution ready'),
    ]
    add_styled_table(doc,
        ['Channel', 'Phase', 'Description'],
        channels)

    doc.add_page_break()

    # =====================================================================
    # APPENDIX C: REFERENCE REPO LOOKUP
    # =====================================================================
    doc.add_heading('Appendix C: Reference Repo Quick Lookup', level=1)

    doc.add_paragraph('All repos cloned in ~/Downloads/external repos/')

    refs = [
        ('Complexity scoring (0-14)', 'DevTeam', 'agents/orchestration/task-loop.md'),
        ('Task loop iterations', 'DevTeam', 'agents/orchestration/task-loop.md'),
        ('Quality gate commands', 'DevTeam', 'agents/orchestration/quality-gate-enforcer.md'),
        ('Model escalation config', 'DevTeam', '.devteam/task-loop-config.yaml'),
        ('Scope validation with VETO', 'DevTeam', 'agents/orchestration/scope-validator.md'),
        ('Anti-abandonment regex', 'DevTeam', 'hooks/persistence-hook.sh'),
        ('SQLite state schema', 'DevTeam', 'scripts/schema.sql'),
        ('Bug Council (5 analysts)', 'DevTeam', 'agents/diagnosis/*.md'),
        ('Cost tracking tables', 'DevTeam', 'scripts/schema.sql'),
        ('File-based output chain', 'wshobson/agents', 'plugins/full-stack-orchestration/commands/full-stack-feature.md'),
        ('State.json resumption', 'wshobson/agents', 'Same file'),
        ('Progressive skill tiers', 'wshobson/agents', 'docs/agent-skills.md'),
        ('Conductor TDD workflow', 'wshobson/agents', 'plugins/conductor/commands/implement.md'),
        ('Model tier assignments', 'wshobson/agents', 'docs/agents.md'),
        ('Phase checkpoints', 'wshobson/agents', 'plugins/full-stack-orchestration/commands/full-stack-feature.md'),
        ('TOML declarative hooks', 'Multi-Agent Squad', '.claude/hooks/enterprise-workflow.toml'),
        ('Dynamic hook generation', 'Multi-Agent Squad', 'scripts/generate-hooks.py'),
        ('Deep agent personas', 'Multi-Agent Squad', '.claude/agents/*/*.md'),
        ('Git worktree orchestration', 'Multi-Agent Squad', 'scripts/worktree-manager.sh'),
        ('Human checkpoint pattern', 'Multi-Agent Squad', '.claude/hooks/enterprise-workflow.toml'),
    ]
    add_styled_table(doc,
        ['Pattern', 'Project', 'File Path'],
        refs)

    # =====================================================================
    # SAVE
    # =====================================================================
    output_path = '/Users/eduardo.torres/Downloads/AgentStudio/AgentStudio-Upgrade-Plan.docx'
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    path = build_document()
    print(f'DOCX generated: {path}')
